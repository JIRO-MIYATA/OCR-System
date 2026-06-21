import streamlit as st
import pandas as pd
import docx
import io
import os
import json
import base64
from datetime import datetime
from pathlib import Path
from string import Template
from PIL import Image
from dotenv import load_dotenv
from google import genai
from google.genai import types

# .envファイルから環境変数をロード
load_dotenv()

# ---------------------------------------------------------
# 定数定義
# ---------------------------------------------------------
BASE_DIR = Path(__file__).parent
EXCEL_FILE = str(BASE_DIR / "extracted_data.xlsx")
GEMINI_MODEL = "gemini-2.5-flash"

# ---------------------------------------------------------
# ページ初期設定
# ---------------------------------------------------------
st.set_page_config(
    page_title="AI OCR & データ抽出システム",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ---------------------------------------------------------
# テーマ設定とカスタムCSSの適用
# ---------------------------------------------------------
if "theme" not in st.session_state:
    st.session_state.theme = "light"


def toggle_theme():
    st.session_state.theme = "dark" if st.session_state.theme == "light" else "light"


IS_DARK = st.session_state.theme == "dark"

# CSS変数の定義
if IS_DARK:
    bg_color = "#09090b"
    bg_subtle = "#0c0c0f"
    card_color = "#0c0c0f"
    border_color = "#1e1e24"
    border_subtle = "#16161a"
    text_color = "#fafafa"
    text_muted = "#a1a1aa"
    shadow = "none"
else:
    bg_color = "#ffffff"
    bg_subtle = "#f9fafb"
    card_color = "#ffffff"
    border_color = "#e4e4e7"
    border_subtle = "#f0f0f2"
    text_color = "#09090b"
    text_muted = "#71717a"
    shadow = "0 1px 3px rgba(0,0,0,0.04), 0 1px 2px rgba(0,0,0,0.03)"

# CSSテンプレートを外部ファイルから読み込み、テーマ変数を適用
_css_path = BASE_DIR / "style.css"
_css_template = Template(_css_path.read_text(encoding="utf-8"))
_css_content = _css_template.safe_substitute(
    bg_color=bg_color,
    bg_subtle=bg_subtle,
    card_color=card_color,
    border_color=border_color,
    border_subtle=border_subtle,
    text_color=text_color,
    text_muted=text_muted,
    shadow=shadow,
)
st.markdown(f"<style>{_css_content}</style>", unsafe_allow_html=True)

# ---------------------------------------------------------
# 解析ヘルパー関数
# ---------------------------------------------------------


def extract_text_from_word(file_bytes):
    """Wordファイルからテキストを抽出する"""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        full_text = []
        # 段落の読み込み
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        # 表の読み込み（セル結合による重複テキストを排除）
        for table in doc.tables:
            for row in table.rows:
                seen = set()
                row_text = []
                for cell in row.cells:
                    cell_id = id(cell._tc)
                    if cell_id not in seen and cell.text.strip():
                        seen.add(cell_id)
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
        return "\n".join(full_text)
    except (ValueError, IOError) as e:
        st.error(f"Wordファイルの解析中にエラーが発生しました: {e}")
        return ""
    except Exception as e:
        st.error(f"Wordファイルの解析中に予期しないエラーが発生しました: {e}")
        return ""


def extract_text_from_excel(file_bytes):
    """Excelファイルからテキスト（Markdownの表形式）を抽出する"""
    try:
        xl = pd.ExcelFile(io.BytesIO(file_bytes))
        full_text = []
        for sheet_name in xl.sheet_names:
            df = xl.parse(sheet_name)
            if not df.empty:
                full_text.append(f"### シート: {sheet_name}")
                # 空白の行・列をドロップし、文字列変換
                df_clean = df.dropna(how='all').fillna("")
                markdown_table = df_clean.to_markdown(index=False)
                full_text.append(markdown_table)
        return "\n\n".join(full_text)
    except (ValueError, IOError) as e:
        st.error(f"Excelファイルの解析中にエラーが発生しました: {e}")
        return ""
    except Exception as e:
        st.error(f"Excelファイルの解析中に予期しないエラーが発生しました: {e}")
        return ""


# ---------------------------------------------------------
# Gemini API 連携ロジック
# ---------------------------------------------------------


@st.cache_resource
def _get_gemini_client(api_key):
    """Gemini APIクライアントをキャッシュして再利用する"""
    return genai.Client(api_key=api_key)


def extract_data_with_gemini(api_key, contents, fields):
    """Gemini APIを使用して、テキストまたは画像から構造化されたJSONデータを抽出する"""
    try:
        client = _get_gemini_client(api_key)

        # 動的にJSONスキーマを構築
        properties = {
            field: types.Schema(
                type=types.Type.STRING,
                description=f"ドキュメントから'{field}'に該当する情報を抽出してください。情報がない場合はnullにしてください。"
            )
            for field in fields
        }

        schema = types.Schema(
            type=types.Type.OBJECT,
            properties=properties,
            required=fields
        )

        fields_desc = ", ".join([f'"{f}"' for f in fields])
        prompt = f"""
        提供されたドキュメント（画像またはテキスト）から、以下の項目を正確に抽出してください。
        抽出項目: {fields_desc}

        【抽出ルール】
        1. 金額や日付などは、可能な限りクリーンな形式（金額ならカンマなし数値、日付ならYYYY-MM-DD形式など）に整形して抽出してください。
        2. ドキュメント内に明確な情報がない項目は null にしてください。
        """

        # リクエストコンテンツの構築
        request_contents = []
        if isinstance(contents, list):
            request_contents.extend(contents)
        else:
            request_contents.append(contents)

        request_contents.append(prompt)

        # API呼び出し
        response = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=request_contents,
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                response_schema=schema,
                temperature=0.1
            )
        )

        # 結果のパース
        return json.loads(response.text)

    except json.JSONDecodeError as e:
        raise RuntimeError(f"APIレスポンスのJSON解析に失敗しました: {e}")
    except Exception as e:
        raise RuntimeError(f"Gemini APIによる解析中にエラーが発生しました: {e}")


# ---------------------------------------------------------
# Excel保存・読み込みロジック
# ---------------------------------------------------------


def load_database():
    """蓄積されたExcelデータをロードする"""
    if os.path.exists(EXCEL_FILE):
        try:
            return pd.read_excel(EXCEL_FILE)
        except (ValueError, IOError) as e:
            st.error(f"データベースファイルの読み込みに失敗しました: {e}")
            return pd.DataFrame()
    return pd.DataFrame()


def save_dataframe_to_excel(df):
    """DataFrameをそのままExcelファイルに書き出す"""
    try:
        df.to_excel(EXCEL_FILE, index=False)
        return True
    except (ValueError, IOError) as e:
        st.error(f"データベースファイルの書き込みに失敗しました: {e}")
        return False


def save_records_to_excel(records):
    """抽出した複数のレコードをExcelに一括追記・保存する"""
    if not records:
        return pd.DataFrame()

    df_new = pd.DataFrame(records)
    if os.path.exists(EXCEL_FILE):
        try:
            df_existing = pd.read_excel(EXCEL_FILE)
            df_combined = pd.concat([df_existing, df_new], ignore_index=True)
        except (ValueError, IOError) as e:
            st.warning(f"既存のExcelの読み込みに失敗したため、新規に作成します: {e}")
            df_combined = df_new
    else:
        df_combined = df_new

    save_dataframe_to_excel(df_combined)
    return df_combined





def filter_database_by_date(df, start_date, end_date):
    """指定期間（開始日〜終了日）のデータにフィルタリングする"""
    if df.empty or '処理日時' not in df.columns:
        return pd.DataFrame(), pd.DataFrame()
    
    df_temp = df.copy()
    # 処理日時を一時的に date 型に変換して判定
    df_temp['_temp_date'] = pd.to_datetime(df_temp['処理日時']).dt.date
    
    # 範囲内のマスク
    mask = (df_temp['_temp_date'] >= start_date) & (df_temp['_temp_date'] <= end_date)
    
    # 範囲内（対象）と範囲外（残すデータ）に分離
    df_target = df[mask]
    df_remaining = df[~mask]
    
    return df_target, df_remaining


def delete_records_by_date(start_date, end_date):
    """指定期間のレコードを削除する"""
    df = load_database()
    if df.empty:
        return False, 0
    
    df_target, df_remaining = filter_database_by_date(df, start_date, end_date)
    deleted_count = len(df_target)
    
    if deleted_count == 0:
        return False, 0
        
    success = save_dataframe_to_excel(df_remaining)
    return success, deleted_count


# ---------------------------------------------------------
# UI 構築
# ---------------------------------------------------------

# ヘッダー
head_left, head_right = st.columns([9, 2])
with head_left:
    st.markdown("""
    <div class="brand">
        <span class="brand-title">📄 Smart Document OCR</span>
    </div>
    <div class="brand-subtitle">
        PDF、Word、Excel、画像をAIで解析し、Excelデータベースへ自動蓄積します。
    </div>
    """, unsafe_allow_html=True)
with head_right:
    theme_label = "☀️ Light" if IS_DARK else "🌙 Dark"
    st.button(theme_label, on_click=toggle_theme, use_container_width=True, key="theme_toggle_btn")

# サイドバー設定
st.sidebar.markdown("### ⚙️ システム設定")

# 1. APIキーの設定 (環境変数優先、なければ入力フォーム)
env_api_key = os.getenv("GEMINI_API_KEY")
if env_api_key:
    api_key = env_api_key
    st.sidebar.success("🔑 Gemini APIキーを環境変数からロードしました")
else:
    api_key = st.sidebar.text_input("🔑 Gemini APIキーを入力してください", type="password")
    if not api_key:
        st.sidebar.warning("⚠️ APIキーが設定されていません。")

# 2. 抽出項目の設定
st.sidebar.markdown("#### 📌 抽出する項目")
st.sidebar.info("抽出したい項目名をカンマ区切りで入力してください。")
default_fields = "日付, 取引先, 金額, 品目"
fields_input = st.sidebar.text_area("抽出項目定義", value=default_fields, height=100)
fields = [f.strip() for f in fields_input.split(",") if f.strip()]



# メインパネルのレイアウト
upload_tab, db_tab = st.tabs(["📤 ファイルのアップロードと解析", "📊 蓄積データの一覧"])

# ---------------------------------------------------------
# タブ1: アップロードと解析
# ---------------------------------------------------------
with upload_tab:
    st.markdown('<div class="custom-card">', unsafe_allow_html=True)
    st.subheader("ファイルのアップロード")
    uploaded_files = st.file_uploader(
        "PDF、Word、Excel、JPEG、PNGファイルをアップロードしてください（複数一括可能）",
        type=["pdf", "docx", "xlsx", "xls", "jpg", "jpeg", "png"],
        accept_multiple_files=True
    )
    st.markdown('</div>', unsafe_allow_html=True)

    # session_state の初期化
    if "extraction_results" not in st.session_state:
        st.session_state.extraction_results = []
    if "file_previews" not in st.session_state:
        st.session_state.file_previews = []
    if "save_completed" not in st.session_state:
        st.session_state.save_completed = False

    def _clear_edit_keys():
        """前回の編集ウィジェットのキーを session_state から除去する"""
        for key in list(st.session_state.keys()):
            if key.startswith("edit_") or key.startswith("preview_text_"):
                del st.session_state[key]

    if uploaded_files:
        st.markdown(f"**アップロードされたファイル数: {len(uploaded_files)} 件**")

        if not api_key:
            st.error("❌ Gemini APIキーが設定されていません。サイドバーまたは環境変数で設定してください。")
        else:
            # AI解析実行ボタン
            if st.button("🚀 AIで一括解析を実行", type="primary", key="run_extraction_btn"):
                _clear_edit_keys()
                progress_bar = st.progress(0)
                status_text = st.empty()

                results = []
                previews = []

                for idx, file in enumerate(uploaded_files):
                    status_text.write(f"【{file.name}】を処理中... ({idx+1}/{len(uploaded_files)})")
                    file_ext = os.path.splitext(file.name)[1].lower()
                    file_bytes = file.read()
                    file.seek(0)

                    contents = None
                    preview_info = {
                        "name": file.name,
                        "ext": file_ext,
                        "preview_type": None,
                        "preview_data": None,
                    }

                    try:
                        # ファイル形式に応じたパース処理
                        if file_ext in [".jpg", ".jpeg", ".png"]:
                            contents = [Image.open(io.BytesIO(file_bytes))]
                            preview_info["preview_type"] = "image"
                            preview_info["preview_data"] = file_bytes
                        elif file_ext == ".pdf":
                            contents = types.Part.from_bytes(
                                data=file_bytes,
                                mime_type="application/pdf"
                            )
                            preview_info["preview_type"] = "pdf"
                            preview_info["preview_data"] = file_bytes
                        elif file_ext == ".docx":
                            text = extract_text_from_word(file_bytes)
                            if not text:
                                raise ValueError("Wordファイルからテキストを抽出できませんでした。")
                            contents = text
                            preview_info["preview_type"] = "text"
                            preview_info["preview_data"] = text
                        elif file_ext in [".xlsx", ".xls"]:
                            text = extract_text_from_excel(file_bytes)
                            if not text:
                                raise ValueError("Excelファイルからデータを抽出できませんでした。")
                            contents = text
                            preview_info["preview_type"] = "text"
                            preview_info["preview_data"] = text

                        # Gemini API を呼び出してデータを抽出
                        if contents:
                            extracted_data = extract_data_with_gemini(api_key, contents, fields)
                            results.append({
                                "file_name": file.name,
                                "file_ext": file_ext,
                                "status": "成功",
                                "extracted": extracted_data,
                                "fields": fields[:],
                            })
                        else:
                            results.append({
                                "file_name": file.name,
                                "file_ext": file_ext,
                                "status": "失敗 (解析可能なコンテンツがありません)",
                                "extracted": None,
                                "fields": fields[:],
                            })
                    except Exception as e:
                        st.error(f"エラー: {file.name} の処理中にエラーが発生しました: {e}")
                        results.append({
                            "file_name": file.name,
                            "file_ext": file_ext,
                            "status": f"失敗 ({str(e)})",
                            "extracted": None,
                            "fields": fields[:],
                        })

                    previews.append(preview_info)
                    progress_bar.progress((idx + 1) / len(uploaded_files))

                st.session_state.extraction_results = results
                st.session_state.file_previews = previews
                st.session_state.save_completed = False
                status_text.success("✨ AI解析が完了しました！下記で内容を確認・修正してください。")
                st.rerun()

            # --------------------------------------------------
            # 解析結果の表示（左: プレビュー / 右: 編集フォーム）
            # --------------------------------------------------
            if st.session_state.extraction_results and not st.session_state.save_completed:
                st.markdown("---")
                st.subheader("📋 解析結果の確認・修正")
                st.info("💡 左側でファイル内容を確認し、右側で抽出データを修正できます。確認後「✅ 確認して保存」を押してください。")

                ok_count = sum(1 for r in st.session_state.extraction_results if r["extracted"])
                ng_count = len(st.session_state.extraction_results) - ok_count
                col_s1, col_s2, col_s3 = st.columns(3)
                col_s1.metric("合計", f"{len(st.session_state.extraction_results)} 件")
                col_s2.metric("成功", f"{ok_count} 件")
                col_s3.metric("失敗", f"{ng_count} 件")

                for idx, result in enumerate(st.session_state.extraction_results):
                    is_ok = result["extracted"] is not None
                    icon = "✅" if is_ok else "❌"

                    with st.expander(f"{icon} {result['file_name']}", expanded=is_ok):
                        if not is_ok:
                            st.error(f"AI解析に失敗しました: {result['status']}")
                            continue

                        col_preview, col_edit = st.columns([1, 1])

                        # --- 左カラム: ファイルプレビュー ---
                        with col_preview:
                            st.markdown("**📎 ファイルプレビュー**")
                            preview = (
                                st.session_state.file_previews[idx]
                                if idx < len(st.session_state.file_previews)
                                else None
                            )
                            if preview and preview["preview_type"]:
                                ptype = preview["preview_type"]
                                pdata = preview["preview_data"]

                                if ptype == "image":
                                    st.image(pdata, use_container_width=True)
                                elif ptype == "pdf":
                                    b64 = base64.b64encode(pdata).decode()
                                    iframe_html = (
                                        f'<iframe src="data:application/pdf;base64,{b64}" '
                                        f'width="100%" height="450" '
                                        f'style="border:1px solid {border_color};'
                                        f'border-radius:8px;"></iframe>'
                                    )
                                    st.markdown(iframe_html, unsafe_allow_html=True)
                                elif ptype == "text":
                                    st.text_area(
                                        "抽出テキスト",
                                        pdata,
                                        height=350,
                                        disabled=True,
                                        key=f"preview_text_{idx}",
                                    )
                            else:
                                st.warning("プレビューデータがありません。")

                        # --- 右カラム: 編集可能な抽出データ ---
                        with col_edit:
                            st.markdown("**✏️ 抽出データ（編集可能）**")
                            for field in result["fields"]:
                                val = result["extracted"].get(field, "") or ""
                                st.text_input(
                                    field,
                                    value=str(val),
                                    key=f"edit_{idx}_{field}",
                                )

                # 保存 / 破棄 ボタン
                st.markdown("---")
                col_discard, col_spacer, col_save = st.columns([1, 2, 1])
                with col_discard:
                    if st.button("🗑️ 破棄する", use_container_width=True, key="discard_btn"):
                        st.session_state.extraction_results = []
                        st.session_state.file_previews = []
                        _clear_edit_keys()
                        st.rerun()
                with col_save:
                    if st.button("✅ 確認して保存", type="primary", use_container_width=True, key="save_btn"):
                        records = []
                        for i, res in enumerate(st.session_state.extraction_results):
                            if res["extracted"] is None:
                                continue
                            edited = {
                                f: st.session_state.get(f"edit_{i}_{f}", "")
                                for f in res["fields"]
                            }
                            records.append({
                                "処理日時": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                                "元ファイル名": res["file_name"],
                                "ファイル形式": res["file_ext"].upper().replace(".", ""),
                                **edited,
                            })

                        if records:
                            save_records_to_excel(records)
                            st.session_state.save_completed = True
                            st.session_state.extraction_results = []
                            st.session_state.file_previews = []
                            _clear_edit_keys()
                            st.rerun()
                        else:
                            st.warning("保存可能なデータがありません。")

            # 保存完了メッセージ
            if st.session_state.get("save_completed"):
                st.success("✨ データの保存が完了しました！「蓄積データの一覧」タブで確認できます。")
                st.session_state.save_completed = False
    else:
        # ファイルがクリアされた場合は結果もリセット
        if st.session_state.get("extraction_results"):
            st.session_state.extraction_results = []
            st.session_state.file_previews = []
            _clear_edit_keys()

# ---------------------------------------------------------
# タブ2: 蓄積データの一覧 & データ管理
# ---------------------------------------------------------
with db_tab:
    df_db = load_database()

    if df_db.empty:
        st.info("📂 まだ蓄積されたデータはありません。ファイルを解析するとここに追加されます。")
    else:
        # 1. 全蓄積データ一覧
        st.markdown('<div class="custom-card">', unsafe_allow_html=True)
        st.subheader("🗂️ 全蓄積データ一覧")

        # カラム順序を整理 (処理日時、ファイル名、ファイル形式を左側に配置)
        meta_cols = ["処理日時", "元ファイル名", "ファイル形式"]
        other_cols = [col for col in df_db.columns if col not in meta_cols]
        final_cols = [col for col in meta_cols if col in df_db.columns] + other_cols
        df_display = df_db[final_cols]

        # Streamlit標準のインタラクティブテーブルで表示
        st.dataframe(df_display, use_container_width=True)

        # Excelファイルのダウンロードボタン
        with open(EXCEL_FILE, "rb") as f:
            st.download_button(
                label="📥 すべてのデータをExcelでダウンロード",
                data=f,
                file_name="ocr_extracted_all_database.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
                key="dl_all_btn"
            )
        st.markdown('</div>', unsafe_allow_html=True)

        # 2. 期間指定の操作セクション
        st.markdown('<div class="custom-card">', unsafe_allow_html=True)
        st.subheader("⚙️ データの期間指定操作（抽出・削除）")
        st.write("ファイルの処理日時を基準にして、データの抽出や削除を行います。")

        # 日付選択의 最小値と最大値を取得
        try:
            db_dates = pd.to_datetime(df_db['処理日時']).dt.date
            min_date = db_dates.min()
            max_date = db_dates.max()
        except Exception:
            min_date = datetime.now().date()
            max_date = datetime.now().date()

        col_date1, col_date2 = st.columns(2)
        with col_date1:
            start_date = st.date_input("開始日", min_date, key="op_start_date")
        with col_date2:
            end_date = st.date_input("終了日", max_date, key="op_end_date")

        if start_date > end_date:
            st.error("❌ エラー: 開始日は終了日より前の日付を指定してください。")
        else:
            df_target, _ = filter_database_by_date(df_db, start_date, end_date)
            
            # メタカラム順序の整理（表示用）
            if not df_target.empty:
                df_target_display = df_target[final_cols]
            else:
                df_target_display = df_target

            # 2a. 抽出
            with st.expander("📥 期間指定データ抽出（プレビューとダウンロード）", expanded=True):
                if df_target.empty:
                    st.info("指定された期間に該当するデータはありません。")
                else:
                    st.write(f"**指定期間のデータ: {len(df_target)} 件**")
                    st.dataframe(df_target_display, use_container_width=True)

                    col_fmt, col_btn = st.columns([1, 1])
                    with col_fmt:
                        export_format = st.radio("ダウンロード形式", ["Excel", "CSV"], horizontal=True, key="export_fmt")
                    with col_btn:
                        st.write("") # ボタンの縦位置調整用のダミー
                        if export_format == "Excel":
                            excel_buffer = io.BytesIO()
                            with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                                df_target.to_excel(writer, index=False)
                            excel_data = excel_buffer.getvalue()

                            st.download_button(
                                label="📥 選択データをExcelで保存",
                                data=excel_data,
                                file_name=f"ocr_extracted_{start_date}_{end_date}.xlsx",
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                use_container_width=True,
                                key="dl_target_excel"
                            )
                        else:
                            csv_data = df_target.to_csv(index=False).encode('utf-8')
                            st.download_button(
                                label="📥 選択データをCSVで保存",
                                data=csv_data,
                                file_name=f"ocr_extracted_{start_date}_{end_date}.csv",
                                mime="text/csv",
                                use_container_width=True,
                                key="dl_target_csv"
                            )

            # 2b. 削除
            with st.expander("🗑️ 期間指定データ削除（プレビューと実行）", expanded=False):
                if df_target.empty:
                    st.info("指定された期間に該当するデータはありません。")
                else:
                    st.warning(f"⚠️ 指定期間のデータ {len(df_target)} 件が削除されます。削除されたデータは復元できません。")
                    st.dataframe(df_target_display, use_container_width=True)

                    # 二重確認用チェックボックス
                    confirm_delete = st.checkbox("はい、表示されているデータを完全に削除することに同意します。", key="confirm_del")

                    if st.button("🚨 選択したデータを完全に削除する", type="primary", disabled=not confirm_delete, use_container_width=True):
                        success, deleted_count = delete_records_by_date(start_date, end_date)
                        if success:
                            st.success(f"✨ {deleted_count} 件のデータを正常に削除しました。")
                            # 画面のリロード
                            st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)
